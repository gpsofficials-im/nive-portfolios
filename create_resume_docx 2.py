from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('NIVETHA G')
title_run.font.size = Pt(24)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(45, 52, 54)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('Full Stack Developer | Computer Science Student')
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor(108, 92, 231)

# Contact Info
contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_run = contact.add_run('Thirupathur, TN, India | +91 6380431813 | nivethanivetha2109@gmail.com')
contact_run.font.size = Pt(10)

doc.add_paragraph()  # Spacing

# ABOUT ME
about_heading = doc.add_paragraph()
about_heading_run = about_heading.add_run('ABOUT ME')
about_heading_run.font.size = Pt(12)
about_heading_run.font.bold = True
about_heading_run.font.color.rgb = RGBColor(45, 52, 54)

about_text = doc.add_paragraph(
    'Enthusiastic Computer Science student at Thiruvalluvar University with strong foundation in programming and web development. '
    'Passionate about creating efficient, user-friendly applications using full-stack technologies. Quick learner with proven ability to adapt to new environments and solve complex problems.'
)
about_text.paragraph_format.space_after = Pt(6)

# EDUCATION
edu_heading = doc.add_paragraph()
edu_heading_run = edu_heading.add_run('EDUCATION')
edu_heading_run.font.size = Pt(12)
edu_heading_run.font.bold = True
edu_heading_run.font.color.rgb = RGBColor(45, 52, 54)

edu_item = doc.add_paragraph()
edu_item_run = edu_item.add_run('Bachelor\'s Degree in Computer Science')
edu_item_run.font.bold = True
edu_item_run = edu_item.add_run(' | Thiruvalluvar University')
edu_item_run.font.italic = True
edu_item.add_run('\nJul 2023 – Present').font.size = Pt(9)
edu_item.paragraph_format.space_after = Pt(6)

coursework = doc.add_paragraph('Coursework: Data Structures, Web Development, Database Systems, Software Engineering', style='List Bullet')
coursework.paragraph_format.space_after = Pt(6)

# EXPERIENCE
exp_heading = doc.add_paragraph()
exp_heading_run = exp_heading.add_run('WORK EXPERIENCE')
exp_heading_run.font.size = Pt(12)
exp_heading_run.font.bold = True
exp_heading_run.font.color.rgb = RGBColor(45, 52, 54)

exp_item = doc.add_paragraph()
exp_item_run = exp_item.add_run('Full Stack Developer Intern')
exp_item_run.font.bold = True
exp_item_run = exp_item.add_run(' | Edu Tantr')
exp_item_run.font.italic = True
exp_item.add_run('\nMay 2025 – Jun 2025').font.size = Pt(9)
exp_item.paragraph_format.space_after = Pt(3)

experiences = [
    'Designed and implemented responsive web interfaces using HTML, CSS, and JavaScript',
    'Developed RESTful APIs using Express.js and managed data with MySQL and MongoDB',
    'Participated in code reviews and debugging, contributing to 65% bug reduction',
    'Optimized application load times and enhanced overall user experience'
]

for exp in experiences:
    p = doc.add_paragraph(exp, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)

# SKILLS
skills_heading = doc.add_paragraph()
skills_heading_run = skills_heading.add_run('SKILLS')
skills_heading_run.font.size = Pt(12)
skills_heading_run.font.bold = True
skills_heading_run.font.color.rgb = RGBColor(45, 52, 54)

skills_data = [
    ('Programming Languages:', 'C++, Python, R, JavaScript'),
    ('Web Technologies:', 'HTML, CSS, JavaScript, Express.js'),
    ('Databases:', 'MySQL, MongoDB, DBMS'),
    ('Tools & OS:', 'Windows 11, VS Code, Git')
]

for skill_category, skill_list in skills_data:
    p = doc.add_paragraph()
    p_run = p.add_run(skill_category)
    p_run.font.bold = True
    p.add_run(f' {skill_list}').font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

# CERTIFICATIONS
cert_heading = doc.add_paragraph()
cert_heading_run = cert_heading.add_run('CERTIFICATIONS')
cert_heading_run.font.size = Pt(12)
cert_heading_run.font.bold = True
cert_heading_run.font.color.rgb = RGBColor(45, 52, 54)

certifications = [
    ('Internship Full Stack Development Certificate', 'Edu Tantr', 'May 2025'),
    ('Oracle Certified Foundation Associate', 'Oracle University', 'Aug 2025'),
    ('Basics of Python', 'Infosys', 'Mar 2025'),
    ('Explore Machine Learning Using Python', 'Infosys', 'Mar 2025'),
    ('Object Oriented Programming Using Python', 'Infosys', 'Mar 2025'),
    ('Introduction to MongoDB for Students', 'ICT Academy', 'Aug 2024')
]

for cert_name, issuer, date in certifications:
    p = doc.add_paragraph()
    p_run = p.add_run(f'• {cert_name}')
    p_run.font.size = Pt(10)
    p.add_run(f' | {issuer} ({date})').font.size = Pt(9)
    p.paragraph_format.space_after = Pt(2)

# PROJECTS
proj_heading = doc.add_paragraph()
proj_heading_run = proj_heading.add_run('PROJECTS')
proj_heading_run.font.size = Pt(12)
proj_heading_run.font.bold = True
proj_heading_run.font.color.rgb = RGBColor(45, 52, 54)

proj_item = doc.add_paragraph()
proj_item_run = proj_item.add_run('Design and Implementation of Human Fitness Tracking System Using CNN Algorithm')
proj_item_run.font.bold = True
proj_item.paragraph_format.space_after = Pt(3)

proj_desc = doc.add_paragraph(
    'Developed a fitness tracking system leveraging Convolutional Neural Networks (CNN) for human activity recognition. '
    'Combines machine learning with practical health tracking applications.'
)
proj_desc.paragraph_format.space_after = Pt(6)

tech_p = doc.add_paragraph()
tech_run = tech_p.add_run('Technologies: ')
tech_run.font.bold = True
tech_p.add_run('Python, CNN, Machine Learning, Deep Learning')
tech_p.paragraph_format.space_after = Pt(12)

# Save the document
doc.save('e:\\NIVE\\NIVETHA G.docx')
print("✓ Resume DOCX created successfully: NIVETHA G.docx")
